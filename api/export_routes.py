import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, PlainTextResponse
from sqlalchemy.ext.asyncio import AsyncSession

from database import get_db
from models import Document

router = APIRouter(tags=["export"])


def extract_plain_text(content: dict) -> str:
    if not content:
        return ""
    lines = []
    def walk(node):
        if isinstance(node, dict):
            if node.get("type") == "text":
                lines.append(node.get("text", ""))
            if node.get("raw"):
                return
            for child in node.get("content", []):
                walk(child)
            if node.get("type") in ("paragraph", "scene_heading", "action",
                                     "character", "dialogue", "heading"):
                lines.append("\n")
        elif isinstance(node, list):
            for item in node:
                walk(item)
    # Handle our simple {raw: "...", nodes: [...]} format
    if content.get("raw"):
        return content["raw"]
    walk(content)
    return "".join(lines).strip()


def content_to_fountain(content: dict, title: str) -> str:
    lines = [f"Title: {title}", "Credit: Written by", "Author: ", "", ""]
    raw = extract_plain_text(content)
    for line in raw.split("\n"):
        stripped = line.strip()
        if not stripped:
            lines.append("")
        elif stripped.startswith(("INT.", "EXT.", "INT./EXT.")):
            lines.append(stripped.upper())
            lines.append("")
        elif stripped.endswith(" TO:") or stripped in ("CUT TO:", "FADE OUT.", "FADE IN:"):
            lines.append("> " + stripped)
            lines.append("")
        else:
            lines.append(stripped)
    return "\n".join(lines)


@router.get("/documents/{doc_id}/export/txt")
async def export_txt(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Not found")
    return PlainTextResponse(
        content=extract_plain_text(doc.content),
        headers={"Content-Disposition": f'attachment; filename="{doc.title}.txt"'}
    )


@router.get("/documents/{doc_id}/export/fountain")
async def export_fountain(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Not found")
    fountain = content_to_fountain(doc.content or {}, doc.title or "Untitled")
    return PlainTextResponse(
        content=fountain,
        headers={"Content-Disposition": f'attachment; filename="{doc.title}.fountain"'}
    )


@router.get("/documents/{doc_id}/export/pdf")
async def export_pdf(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Not found")
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(500, "fpdf2 not installed")
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()
    pdf.set_margins(25, 25, 25)
    pdf.set_font("Courier", style="B", size=14)
    pdf.cell(0, 10, doc.title or "Untitled", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Courier", size=11)
    for line in extract_plain_text(doc.content).split("\n"):
        pdf.multi_cell(0, 6, line or " ")
    pdf.set_font("Courier", style="I", size=8)
    pdf.set_text_color(120, 80, 40)
    pdf.cell(0, 8, "Created with PULP", align="C")
    buf = io.BytesIO(pdf.output())
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc.title}.pdf"'}
    )


@router.get("/documents/{doc_id}/export/docx")
async def export_docx(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Not found")
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed")
    docx = DocxDocument()
    docx.core_properties.title = doc.title or "Untitled"
    title_para = docx.add_heading(doc.title or "Untitled", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in extract_plain_text(doc.content).split("\n"):
        if line.strip():
            docx.add_paragraph(line)
    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc.title}.docx"'}
    )
