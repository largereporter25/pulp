import io
import json
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from api.database import get_db
from api.models import Document

router = APIRouter(tags=["export"])


# The fpdf2 core fonts (Courier/Helvetica/Times) are Latin-1 only, so smart
# punctuation (em dashes, curly quotes, ellipses) must be normalized before
# rendering. This keeps PDF export robust without bundling a Unicode TTF.
_PDF_REPLACEMENTS = {
    "\u2014": "--",  # em dash
    "\u2013": "-",   # en dash
    "\u2018": "'", "\u2019": "'",  # curly single quotes
    "\u201c": '"', "\u201d": '"',  # curly double quotes
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",    # non-breaking space
    "\u2022": "*",    # bullet
}


def pdf_safe(text: str) -> str:
    """Normalize text to a Latin-1-safe form for fpdf2 core fonts."""
    for src, dst in _PDF_REPLACEMENTS.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", "replace").decode("latin-1")


def extract_text_from_content(content_str: str) -> str:
    """Extract plain text from ProseMirror JSON or raw text."""
    if not content_str:
        return ""
    try:
        doc = json.loads(content_str)
        texts = []

        def walk(node):
            if node.get("type") == "text":
                texts.append(node.get("text", ""))
            for child in node.get("content", []):
                walk(child)
            if node.get("type") in ("paragraph", "heading", "blockquote",
                                     "scene_heading", "action", "dialogue",
                                     "character", "parenthetical", "transition"):
                texts.append("\n")

        walk(doc)
        return "".join(texts).strip()
    except Exception:
        return content_str


@router.get("/documents/{doc_id}/export/txt")
async def export_txt(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Document not found")
    text = extract_text_from_content(doc.content)
    filename = doc.title.replace(" ", "_") + ".txt"
    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/documents/{doc_id}/export/fountain")
async def export_fountain(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Document not found")

    lines = [f"Title: {doc.title}", "Credit: Written by", "Draft date: ", "", ""]
    try:
        pm = json.loads(doc.content)

        def walk(node):
            t = node.get("type", "")
            text = "".join(
                c.get("text", "") for c in node.get("content", []) if c.get("type") == "text"
            )
            if t == "scene_heading":
                lines.append(text.upper())
                lines.append("")
            elif t == "action":
                lines.append(text)
                lines.append("")
            elif t == "character":
                lines.append(text.upper())
            elif t == "dialogue":
                lines.append(text)
                lines.append("")
            elif t == "parenthetical":
                lines.append(f"({text})")
            elif t == "transition":
                lines.append(f"> {text.upper()}")
                lines.append("")
            else:
                for child in node.get("content", []):
                    walk(child)

        for child in pm.get("content", []):
            walk(child)
    except Exception:
        lines.append(extract_text_from_content(doc.content))

    fountain_text = "\n".join(lines)
    filename = doc.title.replace(" ", "_") + ".fountain"
    return StreamingResponse(
        io.BytesIO(fountain_text.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/documents/{doc_id}/export/pdf")
async def export_pdf(doc_id: str, db: AsyncSession = Depends(get_db)):
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(500, "fpdf2 not installed. Run: pip install fpdf2")

    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Document not found")

    pdf = FPDF()
    pdf.set_margins(25.4, 25.4, 25.4)  # 1 inch margins
    pdf.add_page()

    # Title page
    pdf.set_font("Courier", "B", 16)
    pdf.set_y(80)
    pdf.cell(0, 10, pdf_safe(doc.title.upper()), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Courier", "", 12)
    pdf.cell(0, 8, "Written by", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(60)
    pdf.set_font("Courier", "", 9)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 6, pdf_safe("Created with PULP — pulp.to"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    pdf.add_page()
    pdf.set_font("Courier", "", 12)

    try:
        pm = json.loads(doc.content)

        def walk(node):
            t = node.get("type", "")
            text = "".join(
                c.get("text", "") for c in node.get("content", []) if c.get("type") == "text"
            )

            text = pdf_safe(text)
            if t == "scene_heading":
                pdf.set_font("Courier", "B", 12)
                pdf.cell(0, 6, text.upper(), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
                pdf.set_font("Courier", "", 12)
            elif t == "action":
                pdf.set_x(25.4)
                pdf.multi_cell(0, 6, text)
                pdf.ln(2)
            elif t == "character":
                pdf.set_x(88.9)  # 3.5 inches from left
                pdf.cell(0, 6, text.upper(), new_x="LMARGIN", new_y="NEXT")
            elif t == "parenthetical":
                pdf.set_x(63.5)  # 2.5 inches
                pdf.cell(0, 6, f"({text})", new_x="LMARGIN", new_y="NEXT")
            elif t == "dialogue":
                pdf.set_x(38.1)  # 1.5 inches, width ~9cm
                pdf.multi_cell(101.6, 6, text)
                pdf.ln(2)
            elif t == "transition":
                pdf.set_font("Courier", "", 12)
                pdf.cell(0, 6, text.upper(), align="R", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            elif t == "paragraph":
                pdf.multi_cell(0, 6, text)
                pdf.ln(2)
            else:
                for child in node.get("content", []):
                    walk(child)

        for child in pm.get("content", []):
            walk(child)
    except Exception:
        plain = pdf_safe(extract_text_from_content(doc.content))
        pdf.multi_cell(0, 6, plain)

    # Footer on each page
    pdf.set_font("Courier", "", 8)
    pdf.set_text_color(150, 150, 150)

    pdf_bytes = pdf.output()
    filename = doc.title.replace(" ", "_") + ".pdf"
    return StreamingResponse(
        io.BytesIO(bytes(pdf_bytes)),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/documents/{doc_id}/export/docx")
async def export_docx(doc_id: str, db: AsyncSession = Depends(get_db)):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")

    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(404, "Document not found")

    d = DocxDocument()
    d.add_heading(doc.title, 0)

    try:
        pm = json.loads(doc.content)

        def walk(node):
            t = node.get("type", "")
            text = "".join(
                c.get("text", "") for c in node.get("content", []) if c.get("type") == "text"
            )
            if t == "scene_heading":
                p = d.add_paragraph(text.upper())
                p.style.font.bold = True
                p.style.font.name = "Courier New"
            elif t in ("action", "paragraph"):
                d.add_paragraph(text)
            elif t == "character":
                p = d.add_paragraph(text.upper())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "dialogue":
                p = d.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "transition":
                p = d.add_paragraph(text.upper())
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif t == "heading":
                level = node.get("attrs", {}).get("level", 1)
                d.add_heading(text, level)
            else:
                for child in node.get("content", []):
                    walk(child)

        for child in pm.get("content", []):
            walk(child)
    except Exception:
        d.add_paragraph(extract_text_from_content(doc.content))

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    filename = doc.title.replace(" ", "_") + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
